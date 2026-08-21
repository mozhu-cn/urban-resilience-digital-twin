"""Repair manuscript.docx for IJDRR submission:
1. Re-insert the corresponding-author email (lost after regeneration)
2. Insert Algorithm 1 (dropped by pandoc) into Section 3.6
"""
import docx
from docx.shared import Pt
from docx.text.paragraph import Paragraph

PATH = r"paper\manuscript.docx"
d = docx.Document(PATH)

# ---------- 1. corresponding author email after 'Independent Researcher, China' ----------
email_text = "*Corresponding author: Zheyu Huang (shuimoqingzhu-cn@qq.com)"
texts = [p.text for p in d.paragraphs]
if not any(email_text in t for t in texts):
    anchor = None
    for p in d.paragraphs:
        if p.text.strip().startswith("Independent Researcher"):
            anchor = p
            break
    assert anchor is not None, "affiliation paragraph not found"
    new_p = docx.oxml.shared.OxmlElement("w:p")
    anchor._p.addnext(new_p)
    para = Paragraph(new_p, anchor._parent)
    run = para.add_run(email_text)
    run.italic = True
    run.font.size = Pt(10)
    print("email inserted after:", anchor.text[:40])
else:
    print("email already present")

# ---------- 2. Algorithm 1 in Section 3.6 ----------
alg_title = "Algorithm 1: High-resolution urban resilience digital twin simulation"
alg_lines = [
    "1.  Initialize terrain elevation, drainage parameters, and infrastructure networks",
    "2.  For each simulation timestep t:",
    "      (a) Update rainfall forcing",
    "      (b) Calculate CA-based flood propagation (Eq. 12)",
    "      (c) Update drainage storage saturation (Eq. 11)",
    "      (d) Evaluate power infrastructure flooding damage (Eq. 14)",
    "      (e) Update communication battery depletion (Eq. 15)",
    "      (f) Calculate system resilience indicator \u03a6(t)",
    "      (g) If failed infrastructure exists: compute restoration priorities (Eq. 17),",
    "          update dynamic road accessibility (Eq. 16), generate optimal repair routes",
    "          (Eq. 18), and execute restoration actions",
    "3.  Output flood evolution and resilience trajectories",
]

if not any(alg_title in p.text for p in d.paragraphs):
    # find the anchor: the paragraph right after "3.6 Simulation Workflow" heading
    anchor = None
    for i, p in enumerate(d.paragraphs):
        if "Simulation Workflow" in p.text and ("3.6" in p.text or i > 0):
            # insert after the paragraph that mentions the simulation process summary
            anchor = p
            break
    # better anchor: paragraph containing 'complete simulation process is summarized'
    for p in d.paragraphs:
        if "summarized in Algorithm" in p.text or "Algorithm" in p.text and "simulation" in p.text:
            anchor = p
            break
    assert anchor is not None, "algorithm anchor not found"
    print("algorithm anchor:", anchor.text[:60])

    def insert_after(par, text, bold=False, italic=False, size=10, indent=False):
        el = docx.oxml.shared.OxmlElement("w:p")
        par._p.addnext(el)
        np_ = Paragraph(el, par._parent)
        run = np_.add_run(text)
        run.bold = bold
        run.italic = italic
        run.font.size = Pt(size)
        if indent:
            pf = np_.paragraph_format
            pf.left_indent = Pt(18)
        return np_

    cur = insert_after(anchor, alg_title, bold=True, size=10)
    for line in alg_lines:
        cur = insert_after(cur, line, size=10, indent=line.startswith("      "))
    print("algorithm inserted")
else:
    print("algorithm already present")

d.save(PATH)

# ---------- verify ----------
d2 = docx.Document(PATH)
full = "\n".join(p.text for p in d2.paragraphs)
print("email ok:", "shuimoqingzhu-cn@qq.com" in full)
print("algorithm ok:", alg_title in full)
print("paragraphs:", len(d2.paragraphs))
