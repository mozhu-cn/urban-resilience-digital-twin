"""Generate supplementary_materials.docx for the IJDRR submission."""
import os
import sys

sys.path.insert(0, os.path.join(os.path.dirname(os.path.abspath(__file__)), "vendor"))

from docx import Document  # noqa: E402
from docx.shared import Pt  # noqa: E402

OUT = os.path.join(os.path.dirname(os.path.abspath(__file__)), "..", "paper",
                   "supplementary_materials.docx")

doc = Document()
style = doc.styles["Normal"]
style.font.name = "Times New Roman"
style.font.size = Pt(11)

doc.add_heading("Supplementary Material", level=0)
doc.add_paragraph(
    "Title: A Physics-Guided Digital Twin Framework for Urban Infrastructure "
    "Resilience Assessment under Extreme Flooding: Integrating "
    "Terrain-Informed Flood Simulation, Delayed Cascading Failures, and "
    "Adaptive Restoration")
doc.add_paragraph("Authors: Zheyu Huang, Yujia Huang")
doc.add_paragraph("Journal: International Journal of Disaster Risk Reduction")

doc.add_heading("S1. Reproducible code and data repository", level=1)
doc.add_paragraph(
    "The complete simulation framework and all experimental outputs are "
    "publicly available in the GitHub repository: "
    "https://github.com/mozhu-cn/urban-resilience-digital-twin")
p = doc.add_paragraph()
r = p.add_run(
    "Repository contents: (1) the full Python implementation of the digital "
    "twin framework, including the terrain-informed cellular-automata flood "
    "model, the interdependent power-communication network model with "
    "time-delayed cascading failures, the adaptive restoration optimizer, and "
    "the visualization modules; (2) configuration files and cached input "
    "data; (3) all experimental scripts (main scenario, two baseline "
    "scenarios, and one-at-a-time sensitivity analysis); (4) all simulation "
    "outputs (frame-wise trajectories, summary metrics) and the "
    "figure-generation scripts; (5) the manuscript sources (LaTeX).")
r.font.size = Pt(11)

doc.add_heading("S2. Input data sources", level=1)
doc.add_paragraph(
    "Road network of Miyazaki City (Japan): OpenStreetMap "
    "(https://www.openstreetmap.org), distributed under the Open Database "
    "License (ODbL). Terrain elevation: Open-Meteo elevation API "
    "(https://open-meteo.com). No proprietary or confidential data were used.")

doc.add_heading("S3. Reproducibility instructions", level=1)
doc.add_paragraph("Requirements: Python 3.10+; packages listed in "
                  "DigitalTwin/requirements.txt (numpy, scipy, pandas, "
                  "networkx, numba, osmnx, plotly, requests, matplotlib).")
doc.add_paragraph(
    "Step 1 - reproduce the complete experiment pipeline (baselines, "
    "sensitivity analysis, and all figures of the manuscript):")
p = doc.add_paragraph()
p.paragraph_format.left_indent = Pt(24)
p.add_run("python run.py --experiments").font.name = "Consolas"
doc.add_paragraph(
    "Step 2 - run the main simulation with the interactive 4D visualization:")
p = doc.add_paragraph()
p.paragraph_format.left_indent = Pt(24)
p.add_run("python run.py").font.name = "Consolas"
doc.add_paragraph(
    "The cached input data are included in the repository, so the pipeline "
    "runs offline. All results are written under results/ and can be compared "
    "with the tables and figures reported in the manuscript.")

doc.add_heading("S4. Correspondence with the manuscript", level=1)
doc.add_paragraph(
    "Fig. 1 corresponds to figures/fig1_study_area.png; Fig. 2 to "
    "fig2_flood_evolution.png; Fig. 3 to fig3_resilience_curves.png; and "
    "Fig. 4 to fig4_sensitivity.png (all 300 dpi). Tables 1-3 of the "
    "manuscript are reproduced by results/baselines_summary.csv and "
    "results/sensitivity_summary.csv.")

doc.save(OUT)
print("saved:", OUT)
